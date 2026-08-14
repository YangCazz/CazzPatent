"""
Convert patent disclosure Markdown to Word (.docx) document.

Integrates CJK typography best practices from minimax-docx:
  - Font pairing: 黑体(SimHei) headings + 宋体(SimSun) body
  - Standard CJK sizes: 三号(16pt) title, 四号(14pt) H1, 小四(12pt) body
  - Line spacing: 1.5× for readability
  - Full inline formatting: **bold**, *italic*, LaTeX protection

Dependencies: pip install python-docx

Usage:
    python md_to_docx.py disclosure.md -o disclosure.docx
"""

import argparse
import json
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# OMML namespace for Word equations
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# ── CJK Typography Constants (from minimax-docx) ────────────────────

# Font sizes in half-points (w:sz)
SIZE_TITLE   = 32  # 三号 16pt — document title
SIZE_H1      = 28  # 四号 14pt — primary heading
SIZE_H2      = 24  # 小四 12pt — secondary heading
SIZE_H3      = 22  # 11pt — tertiary heading
SIZE_BODY    = 24  # 小四 12pt — body text
SIZE_CAPTION = 18  # 9pt — image captions
SIZE_FOOTER  = 18  # 9pt — footnotes

# CJK font names
FONT_HEADING = "黑体"     # SimHei — sans-serif for headings
FONT_BODY    = "宋体"     # SimSun — serif for body
FONT_LATIN   = "Times New Roman"  # fallback for Latin chars
FONT_CAPTION = "宋体"     # SimSun

# Line spacing in 240ths of a line
LINE_SPACING = 360  # 1.5× (360/240)

# Page margins in cm
MARGIN_TOP    = 2.54
MARGIN_BOTTOM = 2.54
MARGIN_LEFT   = 3.18
MARGIN_RIGHT  = 3.18


# ── XML helpers ────────────────────────────────────────────────────

def _set_cjk_font(run, font_heading=False):
    """Set font slots for mixed CJK/Latin rendering on a run."""
    rPr = run._element.get_or_add_rPr()
    latin = FONT_LATIN
    cjk = FONT_HEADING if font_heading else FONT_BODY
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{latin}" w:hAnsi="{latin}" '
        f'w:eastAsia="{cjk}" w:cs="{latin}"/>'
    )
    # Remove existing rFonts if any
    for existing in rPr.findall(qn('w:rFonts')):
        rPr.remove(existing)
    rPr.append(rFonts)


def _set_line_spacing(paragraph, spacing=LINE_SPACING):
    """Set CJK-appropriate line spacing on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    spacing_elem = parse_xml(
        f'<w:spacing {nsdecls("w")} '
        f'w:line="{spacing}" w:lineRule="auto"/>'
    )
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    pPr.append(spacing_elem)


# ── Per-document LaTeX symbol extraction ────────────────────────────

def _extract_latex_commands(md_text: str) -> set:
    """Extract all unique LaTeX command names from math blocks in the document."""
    commands = set()
    # Find all $...$ and $$...$$ blocks
    for m in re.finditer(r"\$\$?(.+?)\$\$?", md_text, re.DOTALL):
        latex = m.group(1)
        # Extract \command sequences (alphabetic commands only)
        for cmd in re.finditer(r"\\([a-zA-Z]+)", latex):
            commands.add(cmd.group(1))
    return commands


def _build_symbol_map(md_text: str) -> dict:
    """Build a document-specific LaTeX→Unicode symbol map.

    Combines the built-in symbol library with any document-specific commands.
    Returns (symbol_map, unknown_commands) tuple.
    """
    used = _extract_latex_commands(md_text)
    sym_map = dict(_SYMBOLS)  # Start with built-in library
    unknown = set()

    for cmd in sorted(used):
        if cmd not in sym_map and cmd not in _SKIP_COMMANDS:
            unknown.add(cmd)

    return sym_map, unknown


# ── LaTeX → OMML (Word equation) converter ─────────────────────────

# Commands that are structural (not symbols to render)
_SKIP_COMMANDS = {
    "frac", "sum", "int", "prod", "mathbb", "mathbf", "boldsymbol",
    "mathcal", "mathit", "bar", "hat", "ddot", "dot", "tilde", "vec",
    "tag", "label", "left", "right", "begin", "end", "\\",
    "text", "operatorname", "mathrm", "sin", "cos", "tan", "log", "ln", "exp",
    "lim", "max", "min", "sup", "inf", "det", "dim", "gcd", "ker",
    "Pr", "argmax", "argmin", "arg",
    "quad", "qquad", "big", "Big", "bigg", "Bigg",
    "underbrace", "overbrace",
}

# Per-document symbol map — set before conversion
_DOC_SYMBOLS: dict = {}
_UNKNOWN_COMMANDS: set = set()

# Greek letters + common math symbols → Unicode
_SYMBOLS = {
    # Greek lowercase
    "alpha": "α", "beta": "β", "gamma": "γ", "delta": "δ",
    "epsilon": "ε", "zeta": "ζ", "eta": "η", "theta": "θ",
    "iota": "ι", "kappa": "κ", "lambda": "λ", "mu": "μ",
    "nu": "ν", "xi": "ξ", "pi": "π", "rho": "ρ",
    "sigma": "σ", "tau": "τ", "upsilon": "υ", "phi": "φ",
    "chi": "χ", "psi": "ψ", "omega": "ω",
    # Greek uppercase
    "Alpha": "Α", "Beta": "Β", "Gamma": "Γ", "Delta": "Δ",
    "Epsilon": "Ε", "Zeta": "Ζ", "Eta": "Η", "Theta": "Θ",
    "Iota": "Ι", "Kappa": "Κ", "Lambda": "Λ", "Mu": "Μ",
    "Nu": "Ν", "Xi": "Ξ", "Pi": "Π", "Rho": "Ρ",
    "Sigma": "Σ", "Tau": "Τ", "Upsilon": "Υ", "Phi": "Φ",
    "Chi": "Χ", "Psi": "Ψ", "Omega": "Ω",
    # Greek variants
    "varepsilon": "ε", "vartheta": "ϑ", "varphi": "φ",
    "varrho": "ϱ", "varsigma": "ς", "digamma": "ϝ",
    # Operators / big symbols
    "infty": "∞", "partial": "∂", "nabla": "∇",
    "times": "×", "cdot": "·", "pm": "±", "mp": "∓",
    "div": "÷", "ast": "∗", "star": "⋆", "circ": "∘",
    "bullet": "∙", "oplus": "⊕", "ominus": "⊖", "otimes": "⊗",
    "oslash": "⊘", "odot": "⊙", "bigcirc": "◯",
    "setminus": "∖", "cap": "∩", "cup": "∪", "sqcap": "⊓", "sqcup": "⊔",
    "wedge": "∧", "vee": "∨", "land": "∧", "lor": "∨", "lnot": "¬", "neg": "¬",
    "triangle": "△", "bigtriangleup": "△", "bigtriangledown": "▽",
    "Box": "□", "Diamond": "◇",
    # Relations
    "leq": "≤", "le": "≤", "geq": "≥", "ge": "≥", "neq": "≠", "ne": "≠", "equiv": "≡",
    "approx": "≈", "sim": "∼", "simeq": "≃", "cong": "≅",
    "propto": "∝", "parallel": "∥", "perp": "⊥",
    "ll": "≪", "gg": "≫", "prec": "≺", "succ": "≻",
    "preceq": "≼", "succeq": "≽", "doteq": "≐",
    # Set relations
    "in": "∈", "notin": "∉", "ni": "∋", "owns": "∋",
    "subset": "⊂", "supset": "⊃", "subseteq": "⊆", "supseteq": "⊇",
    "subsetneq": "⊊", "supsetneq": "⊋",
    "emptyset": "∅", "varnothing": "∅",
    "forall": "∀", "exists": "∃", "nexists": "∄",
    # Arrows
    "rightarrow": "→", "leftarrow": "←", "Rightarrow": "⇒", "Leftarrow": "⇐",
    "leftrightarrow": "↔", "Leftrightarrow": "⇔",
    "mapsto": "↦", "longmapsto": "⟼",
    "uparrow": "↑", "downarrow": "↓",
    "longrightarrow": "⟶", "longleftarrow": "⟵",
    "Longrightarrow": "⟹", "Longleftarrow": "⟸",
    "iff": "⟺",
    "to": "→", "gets": "←",
    "mid": "∣", "nmid": "∤",
    # Various
    "ldots": "…", "dots": "…", "cdots": "⋯", "vdots": "⋮", "ddots": "⋱",
    "angle": "∠", "measuredangle": "∡",
    "aleph": "א", "hbar": "ℏ", "ell": "ℓ", "wp": "℘", "Re": "ℜ", "Im": "ℑ",
    "prime": "′", "backprime": "‵",
    "dag": "†", "ddag": "‡",
    "top": "⊤", "bot": "⊥",
    "iint": "∬", "bigcup": "⋃", "sqrt": "√",
    "Vert": "‖", "lVert": "‖", "rVert": "‖",
    "langle": "⟨", "rangle": "⟩", "lVert": "‖", "rVert": "‖",
    "lfloor": "⌊", "rfloor": "⌋", "lceil": "⌈", "rceil": "⌉",
    # Text operators (rendered as upright text in math)
    # \text{...} and \operatorname{...} are handled separately in _parse_expr
}


# ── 补充：完整符号库（常用数学/物理/统计符号与别名） ──
_SYMBOLS.update({
    # Greek variants (补充)
    "varkappa": "ϰ", "varpi": "ϖ",
    # Binary operators (补充)
    "uplus": "⊎", "amalg": "⨿", "sqsubset": "⊏", "sqsupset": "⊐",
    "sqsubseteq": "⊑", "sqsupseteq": "⊒", "diamond": "⋄", "lozenge": "◊",
    "blacklozenge": "⧫", "triangleleft": "◁", "triangleright": "▷",
    "lhd": "⊲", "rhd": "⊳", "unlhd": "⊴", "unrhd": "⊵",
    "boxplus": "⊞", "boxminus": "⊟", "boxtimes": "⊠", "boxdot": "⊡",
    "circledast": "⊛", "circledcirc": "⊚", "circleddash": "⊝",
    "ltimes": "⋉", "rtimes": "⋊", "leftthreetimes": "⋋", "rightthreetimes": "⋌",
    "curlyvee": "⋎", "curlywedge": "⋏", "intercal": "⊺",
    "dagger": "†", "ddagger": "‡", "Cap": "⋒", "Cup": "⋓",
    "barwedge": "⊼", "veebar": "⊻", "doublebarwedge": "⩞",
    "circledS": "Ⓢ", "smallsetminus": "∖",
    # Relations (补充)
    "lesssim": "≲", "gtrsim": "≳", "lessapprox": "⪅", "gtrapprox": "⪆",
    "lessgtr": "≶", "gtrless": "≷", "lesseqgtr": "⋚", "gtreqless": "⋛",
    "lesseqqgtr": "⪋", "gtreqqless": "⪌", "triangleq": "≜", "approxeq": "≊",
    "circeq": "≗", "bumpeq": "≏", "Bumpeq": "≎", "doteqdot": "≑",
    "fallingdotseq": "≒", "risingdotseq": "≓", "asymp": "≍",
    "bowtie": "⋈", "Join": "⋈", "vdash": "⊢", "dashv": "⊣",
    "models": "⊨", "vDash": "⊨", "Vdash": "⊩", "Vvdash": "⊪",
    "nless": "≮", "ngtr": "≯", "nleq": "≰", "ngeq": "≱",
    "lneq": "⪇", "gneq": "⪈", "lneqq": "≨", "gneqq": "≩",
    "nsim": "≁", "ncong": "≇", "napprox": "≉", "nparallel": "∦",
    "varsubsetneq": "⊊", "varsupsetneq": "⊋", "subsetneqq": "⫋", "supsetneqq": "⫌",
    "nsubseteq": "⊈", "nsupseteq": "⊉", "subseteqq": "⫅", "supseteqq": "⫆",
    "nprec": "⊀", "nsucc": "⊁", "precnsim": "⋨", "succnsim": "⋩",
    "precnapprox": "⪹", "succnapprox": "⪺", "smile": "⌣", "frown": "⌢",
    "therefore": "∴", "because": "∵", "varpropto": "∝",
    "shortmid": "∣", "shortparallel": "∥", "between": "≬", "pitchfork": "⋔",
    "backepsilon": "϶",
    # Set relations (补充)
    "Subset": "⋐", "Supset": "⋑",
    # Arrows (补充)
    "updownarrow": "↕", "Updownarrow": "⇕", "nearrow": "↗", "searrow": "↘",
    "swarrow": "↙", "nwarrow": "↖", "hookrightarrow": "↪", "hookleftarrow": "↩",
    "rightharpoonup": "⇀", "rightharpoondown": "⇁", "leftharpoonup": "↼", "leftharpoondown": "↽",
    "rightleftharpoons": "⇌", "leftrightharpoons": "⇋", "twoheadrightarrow": "↠", "twoheadleftarrow": "↞",
    "dashrightarrow": "⇢", "dashleftarrow": "⇠", "leftrightsquigarrow": "↭",
    "rightsquigarrow": "⇝", "leadsto": "⇝", "curvearrowleft": "↶", "curvearrowright": "↷",
    "circlearrowleft": "↺", "circlearrowright": "↻", "looparrowleft": "↫", "looparrowright": "↬",
    "Lsh": "↰", "Rsh": "↱", "multimap": "⊸", "nrightarrow": "↛", "nleftarrow": "↚",
    "nRightarrow": "⇏", "nLeftarrow": "⇍", "nleftrightarrow": "↮", "upuparrows": "⇈", "downdownarrows": "⇊",
    # Miscellaneous math (补充)
    "iiint": "∭", "iiiint": "⨌", "oint": "∮", "oiint": "∯", "oiiint": "∰",
    "bigcap": "⋂", "bigsqcup": "⨆", "bigvee": "⋁", "bigwedge": "⋀",
    "bigoplus": "⨁", "bigotimes": "⨂", "bigodot": "⨀", "coprod": "∐",
    "blacksquare": "■", "blacktriangle": "▲", "blacktriangledown": "▼",
    "clubsuit": "♣", "diamondsuit": "♢", "heartsuit": "♡", "spadesuit": "♠",
    "flat": "♭", "natural": "♮", "sharp": "♯", "sphericalangle": "∢",
    "ulcorner": "⌜", "urcorner": "⌝", "llcorner": "⌞", "lrcorner": "⌟",
    "backslash": "∖", "eth": "ð", "Finv": "Ⅎ", "Game": "⅁", "Bbbk": "𝕜",
    "diagdown": "╲", "diagup": "╱", "checkmark": "✓", "maltese": "✠",
    # Delimiters (补充)
    "lgroup": "⟮", "rgroup": "⟯",
})

# ── 补充：结构命令（函数名/重音/字体，交给 _parse_expr 处理） ──
_SKIP_COMMANDS.update({
    "mathfrak", "mathsf", "mathtt", "mathscr", "bm",
    "widehat", "widetilde", "check", "breve", "acute", "grave", "mathring",
    "overline", "underline",
    "liminf", "limsup", "arcsin", "arccos", "arctan",
    "sinh", "cosh", "tanh", "coth", "csc", "sec", "cot", "deg", "lg", "hom",
    "sgn", "diag", "tr", "trace", "rank", "span", "erf", "erfc", "adj",
    "injlim", "projlim", "varinjlim", "varprojlim", "varliminf", "varlimsup",
})


def _omml_elem(tag: str, **attrs) -> OxmlElement:
    """Create an OMML element."""
    return OxmlElement(f"m:{tag}", attrs={})


def _omml_run(text: str, normal: bool = False, bold: bool = False) -> OxmlElement:
    """Create an <m:r> with <m:t> text."""
    r = _omml_elem("r")
    if normal or bold:
        rPr = _omml_elem("rPr")
        if normal:
            nor = _omml_elem("nor")
            rPr.append(nor)
        if bold:
            sty = _omml_elem("sty")
            sty.set(qn("m:val"), "b")
            rPr.append(sty)
        r.append(rPr)
    t = _omml_elem("t")
    t.text = text
    r.append(t)
    return r

# Double-struck (blackboard bold) character mapping
_BB = {
    "A": "𝔸", "B": "𝔹", "C": "ℂ", "D": "𝔻", "E": "𝔼",
    "F": "𝔽", "G": "𝔾", "H": "ℍ", "I": "𝕀", "J": "𝕁",
    "K": "𝕂", "L": "𝕃", "M": "𝕄", "N": "ℕ", "O": "𝕆",
    "P": "ℙ", "Q": "ℚ", "R": "ℝ", "S": "𝕊", "T": "𝕋",
    "U": "𝕌", "V": "𝕍", "W": "𝕎", "X": "𝕏", "Y": "𝕐", "Z": "ℤ",
    "a": "𝕒", "b": "𝕓", "c": "𝕔", "d": "𝕕", "e": "𝕖",
    "f": "𝕗", "g": "𝕘", "h": "𝕙", "i": "𝕚", "j": "𝕛",
    "k": "𝕜", "l": "𝕝", "m": "𝕞", "n": "𝕟", "o": "𝕠",
    "p": "𝕡", "q": "𝕢", "r": "𝕣", "s": "𝕤", "t": "𝕥",
    "u": "𝕦", "v": "𝕧", "w": "𝕨", "x": "𝕩", "y": "𝕪", "z": "𝕫",
    "1": "𝟙", "2": "𝟚", "0": "𝟘",
}


def _latex_to_omml(latex: str) -> Optional[OxmlElement]:
    """Convert a LaTeX math expression to OMML.

    Returns an <m:oMath> element on success, None on failure.
    Handles: subscripts, superscripts, fractions, Greek letters,
             sums, integrals, mathbb/mathbf, delimiters.
    """
    latex = latex.strip()
    if not latex:
        return None

    try:
        tokens, _ = _tokenize(latex, 0)
        content = _parse_expr(tokens)
        if content is not None:
            omath = _omml_elem("oMath")
            for child in content:
                omath.append(child)
            return omath
    except Exception:
        pass
    return None


def _parse_expr(tokens: list) -> list | None:
    """Parse a sequence of (token, is_command) tuples into OMML elements."""
    result = []
    i = 0
    while i < len(tokens):
        tok, is_cmd = tokens[i]
        if tok == "^":
            if result and i + 1 < len(tokens):
                base = result.pop()
                sup_expr, consumed = _parse_arg(tokens, i + 1)
                ssup = _omml_elem("sSup")
                e = _omml_elem("e")
                _append_to(e, base)
                ssup.append(e)
                sup_el = _omml_elem("sup")
                for c in sup_expr:
                    _append_to(sup_el, c)
                ssup.append(sup_el)
                result.append(ssup)
                i = consumed + 1
                continue
        elif tok == "_":
            if result and i + 1 < len(tokens):
                base = result.pop()
                sub_expr, consumed = _parse_arg(tokens, i + 1)
                # Check for ^ after the subscript
                if consumed + 1 < len(tokens) and tokens[consumed + 1] == "^":
                    # sub + sup
                    sup_expr, sup_consumed = _parse_arg(tokens, consumed + 2)
                    ssubsup = _omml_elem("sSubSup")
                    e = _omml_elem("e")
                    _append_to(e, base)
                    ssubsup.append(e)
                    sub_el = _omml_elem("sub")
                    for c in sub_expr:
                        _append_to(sub_el, c)
                    ssubsup.append(sub_el)
                    sup_el = _omml_elem("sup")
                    for c in sup_expr:
                        _append_to(sup_el, c)
                    ssubsup.append(sup_el)
                    result.append(ssubsup)
                    i = sup_consumed + 1
                else:
                    ssub = _omml_elem("sSub")
                    e = _omml_elem("e")
                    _append_to(e, base)
                    ssub.append(e)
                    sub_el = _omml_elem("sub")
                    for c in sub_expr:
                        _append_to(sub_el, c)
                    ssub.append(sub_el)
                    result.append(ssub)
                    i = consumed + 1
                continue
        elif tok == "frac" and is_cmd:
            num_expr, consumed1 = _parse_arg(tokens, i + 1)
            den_expr, consumed2 = _parse_arg(tokens, consumed1 + 1)
            f = _omml_elem("f")
            num_el = _omml_elem("num")
            for c in num_expr:
                _append_to(num_el, c)
            f.append(num_el)
            den_el = _omml_elem("den")
            for c in den_expr:
                _append_to(den_el, c)
            f.append(den_el)
            result.append(f)
            i = consumed2 + 1
            continue
        elif tok in ("underbrace", "overbrace") and is_cmd:
            # \underbrace{base}_{annot}  → limLow + groupChr (╯)
            # \overbrace{base}^{annot}   → limUpp + groupChr (╮)
            # Must "eat" the trailing _/^ ourselves: sSub would misplace
            # the annotation to the lower-right instead of centred below.
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            gc = _omml_elem("groupChr")
            gc_pr = _omml_elem("groupChrPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "⏟" if tok == "underbrace" else "⏞")
            gc_pr.append(chr_el)
            pos_el = _omml_elem("pos")
            pos_el.set(qn("m:val"), "bot" if tok == "underbrace" else "top")
            gc_pr.append(pos_el)
            gc.append(gc_pr)
            e = _omml_elem("e")
            for c in arg_expr:
                _append_to(e, c)
            gc.append(e)

            # Check for trailing _{annot} (underbrace) or ^{annot} (overbrace)
            next_i = consumed + 1
            annot_tok = "_" if tok == "underbrace" else "^"
            if next_i < len(tokens) and _tok(tokens, next_i) == annot_tok:
                annot_expr, annot_consumed = _parse_arg(tokens, next_i + 1)
                lim = _omml_elem("limLow" if tok == "underbrace" else "limUpp")
                lim_e = _omml_elem("e")
                lim_e.append(gc)
                lim.append(lim_e)
                lim_annot = _omml_elem("lim")
                for c in annot_expr:
                    _append_to(lim_annot, c)
                lim.append(lim_annot)
                result.append(lim)
                i = annot_consumed + 1
            else:
                result.append(gc)
                i = consumed + 1
            continue
        elif tok == "sum" and is_cmd:
            nary = _omml_elem("nary")
            nary_pr = _omml_elem("naryPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "∑")
            nary_pr.append(chr_el)
            nary.append(nary_pr)
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "_":
                sub_expr, sub_consumed = _parse_arg(tokens, i + 2)
                sub_el = _omml_elem("sub")
                for c in sub_expr:
                    _append_to(sub_el, c)
                nary.append(sub_el)
                i = sub_consumed
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "^":
                sup_expr, sup_consumed = _parse_arg(tokens, i + 2)
                sup_el = _omml_elem("sup")
                for c in sup_expr:
                    _append_to(sup_el, c)
                nary.append(sup_el)
                i = max(i, sup_consumed)
            result.append(nary)
            i += 1
            continue
        elif tok == "int" and is_cmd:
            nary = _omml_elem("nary")
            nary_pr = _omml_elem("naryPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "∫")
            nary_pr.append(chr_el)
            nary.append(nary_pr)
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "_":
                sub_expr, sub_consumed = _parse_arg(tokens, i + 2)
                sub_el = _omml_elem("sub")
                for c in sub_expr:
                    _append_to(sub_el, c)
                nary.append(sub_el)
                i = sub_consumed
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "^":
                sup_expr, sup_consumed = _parse_arg(tokens, i + 2)
                sup_el = _omml_elem("sup")
                for c in sup_expr:
                    _append_to(sup_el, c)
                nary.append(sup_el)
                i = max(i, sup_consumed)
            result.append(nary)
            i += 1
            continue
        elif tok == "prod" and is_cmd:
            nary = _omml_elem("nary")
            nary_pr = _omml_elem("naryPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "∏")
            nary_pr.append(chr_el)
            nary.append(nary_pr)
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "_":
                sub_expr, sub_consumed = _parse_arg(tokens, i + 2)
                sub_el = _omml_elem("sub")
                for c in sub_expr:
                    _append_to(sub_el, c)
                nary.append(sub_el)
                i = sub_consumed
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "^":
                sup_expr, sup_consumed = _parse_arg(tokens, i + 2)
                sup_el = _omml_elem("sup")
                for c in sup_expr:
                    _append_to(sup_el, c)
                nary.append(sup_el)
                i = max(i, sup_consumed)
            result.append(nary)
            i += 1
            continue
        elif tok in ("mathbb", "mathbf", "boldsymbol", "mathcal", "mathit", "mathfrak", "mathsf", "mathtt", "mathscr", "bm") and is_cmd:
            # Styled text commands: \mathbb{R}, \mathbf{v}, \boldsymbol{p}
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            if tok in ("mathbf", "boldsymbol", "bm"):
                # Bold: wrap each run with bold styling
                for c in arg_expr:
                    if c.tag == qn("m:r"):
                        # Replace with bold version
                        text = "".join(t.text or "" for t in c.findall(qn("m:t")))
                        result.append(_omml_run(text, bold=True))
                    else:
                        result.append(c)
            elif tok == "mathbb":
                # Blackboard bold: use Unicode double-struck chars if available
                for c in arg_expr:
                    text = "".join(t.text or "" for t in c.findall(qn("m:t")))
                    mapped = "".join(_BB.get(ch, ch) for ch in text)
                    result.append(_omml_run(mapped))
            else:
                # \mathcal, \mathit: just pass through
                for c in arg_expr:
                    result.append(c)
            i = consumed + 1
            continue
        elif tok == "left" and is_cmd:
            # \left( ... \right)
            left_delim = _tok(tokens, i + 1) if i + 1 < len(tokens) else "("
            d = _omml_elem("d")
            d_pr = _omml_elem("dPr")
            beg = _omml_elem("begChr")
            beg.set(qn("m:val"), _delim_char(left_delim))
            d_pr.append(beg)
            d.append(d_pr)
            e = _omml_elem("e")
            # Find matching \right
            j = i + 2
            depth = 1
            inner_tokens = []
            while j < len(tokens) and depth > 0:
                if _tok(tokens, j) == "left":
                    depth += 1
                elif _tok(tokens, j) == "right":
                    depth -= 1
                    if depth == 0:
                        break
                inner_tokens.append(tokens[j])
                j += 1
            inner = _parse_expr(inner_tokens)
            if inner:
                for c in inner:
                    _append_to(e, c)
            d.append(e)
            # Set end delimiter from the \right's argument
            right_delim = _tok(tokens, j + 1) if j + 1 < len(tokens) else ")"
            end = _omml_elem("endChr")
            end.set(qn("m:val"), _delim_char(right_delim))
            d_pr.append(end)
            result.append(d)
            # Skip past \right and its delimiter
            i = j + 2
            continue
            continue
        elif tok == "right":
            i += 1
            continue
        elif tok == "begin" and is_cmd:
            # \begin{matrix_env} ... \end{matrix_env}
            _MATRIX_DELIMS = {
                "bmatrix": ("[", "]"), "pmatrix": ("(", ")"),
                "vmatrix": ("|", "|"), "Vmatrix": ("‖", "‖"),
                "Bmatrix": ("{", "}"), "matrix": ("", ""),
                "cases": ("{", ""),
            }
            # Read env name: {envname}
            if i + 3 < len(tokens) and _tok(tokens, i + 1) == "{":
                env = _tok(tokens, i + 2)
                if i + 3 < len(tokens) and _tok(tokens, i + 3) == "}" and env in _MATRIX_DELIMS:
                    # Find matching \end{env}
                    depth = 1
                    j = i + 4
                    matrix_tokens = []
                    while j < len(tokens) and depth > 0:
                        t_j, c_j = tokens[j]
                        if c_j and t_j == "begin":
                            depth += 1
                        elif c_j and t_j == "end":
                            # Check if next tokens are {env}
                            if j + 2 < len(tokens) and _tok(tokens, j + 1) == "{" and _tok(tokens, j + 2) == env:
                                depth -= 1
                                if depth == 0:
                                    break
                        matrix_tokens.append(tokens[j])
                        j += 1
                    if depth == 0:
                        # Build OMML matrix from matrix_tokens (split by \\ and &)
                        beg_ch, end_ch = _MATRIX_DELIMS[env]
                        d = _omml_elem("d")
                        if beg_ch or end_ch:
                            d_pr = _omml_elem("dPr")
                            beg = _omml_elem("begChr")
                            beg.set(qn("m:val"), beg_ch)
                            d_pr.append(beg)
                            end = _omml_elem("endChr")
                            end.set(qn("m:val"), end_ch)
                            d_pr.append(end)
                            d.append(d_pr)
                        e_out = _omml_elem("e")
                        m = _omml_elem("m")
                        m_pr = _omml_elem("mPr")
                        m.append(m_pr)
                        # Parse rows
                        rows = _split_matrix_rows(matrix_tokens)
                        ncols = max(len(row) for row in rows) if rows else 0
                        for row_tokens in rows:
                            mr = _omml_elem("mr")
                            for cell_tokens in row_tokens:
                                cell_expr = _parse_expr(cell_tokens)
                                cell_e = _omml_elem("e")
                                for c in cell_expr:
                                    _append_to(cell_e, c)
                                mr.append(cell_e)
                            # Pad to ncols
                            while len(mr) < ncols:
                                cell_e = _omml_elem("e")
                                mr.append(cell_e)
                            m.append(mr)
                        e_out.append(m)
                        d.append(e_out)
                        result.append(d)
                        i = j + 4  # Skip past \end{env}
                        continue
            # Fallback: pass \begin through as text
            result.append(_omml_run("begin"))
            i += 1
            continue
        elif tok == "end" and is_cmd:
            # \end{env} — should have been consumed by begin handler, skip
            i += 1
            continue
        elif tok == "tag":
            _, consumed = _parse_arg(tokens, i + 1)
            i = consumed + 1
            continue
        elif tok == "label":
            _, consumed = _parse_arg(tokens, i + 1)
            i = consumed + 1
            continue
        elif tok in ("text", "operatorname", "mathrm"):
            # \text{SegNet}, \operatorname{argmin}, \mathrm{Var} — upright text in math
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            combined = ""
            for c in arg_expr:
                for t in c.findall(qn("m:t")):
                    if t.text:
                        combined += t.text
            if combined:
                result.append(_omml_run(combined, normal=True))
            i = consumed + 1
            continue
        elif tok in ("quad", "qquad"):
            result.append(_omml_run("  "))
            i += 1
            continue
        elif tok in (";",) and False:
            # Placeholder: \; thick space (currently disabled)
            i += 1
            continue
        elif tok == "bar" and is_cmd:
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            acc = _omml_elem("acc")
            acc_pr = _omml_elem("accPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "̅")
            acc_pr.append(chr_el)
            acc.append(acc_pr)
            e = _omml_elem("e")
            for c in arg_expr:
                _append_to(e, c)
            acc.append(e)
            result.append(acc)
            i = consumed + 1
            continue
        elif tok == "hat":
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            acc = _omml_elem("acc")
            acc_pr = _omml_elem("accPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "̂")
            acc_pr.append(chr_el)
            acc.append(acc_pr)
            e = _omml_elem("e")
            for c in arg_expr:
                _append_to(e, c)
            acc.append(e)
            result.append(acc)
            i = consumed + 1
            continue
        elif tok in ("ddot", "dot", "tilde", "vec"):
            _ACCENTS = {"ddot": "̈", "dot": "̇", "tilde": "̃", "vec": "⃗"}
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            acc = _omml_elem("acc")
            acc_pr = _omml_elem("accPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), _ACCENTS.get(tok, "̇"))
            acc_pr.append(chr_el)
            acc.append(acc_pr)
            e = _omml_elem("e")
            for c in arg_expr:
                _append_to(e, c)
            acc.append(e)
            result.append(acc)
            i = consumed + 1
            continue
        elif tok in ("overline", "underline") and is_cmd:
            # \overline{...} / \underline{...} — wide bar over/under
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            bar_el = _omml_elem("bar")
            bar_pr = _omml_elem("barPr")
            pos_el = _omml_elem("pos")
            pos_el.set(qn("m:val"), "top" if tok == "overline" else "bot")
            bar_pr.append(pos_el)
            bar_el.append(bar_pr)
            e = _omml_elem("e")
            for c in arg_expr:
                _append_to(e, c)
            bar_el.append(e)
            result.append(bar_el)
            i = consumed + 1
            continue
        elif tok in ("widehat", "widetilde", "check", "breve", "acute", "grave", "mathring") and is_cmd:
            _ACCENTS2 = {"widehat": "̂", "widetilde": "̃", "check": "̌", "breve": "̆", "acute": "́", "grave": "̀", "mathring": "̊"}
            arg_expr, consumed = _parse_arg(tokens, i + 1)
            acc = _omml_elem("acc")
            acc_pr = _omml_elem("accPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), _ACCENTS2.get(tok, "̇"))
            acc_pr.append(chr_el)
            acc.append(acc_pr)
            e = _omml_elem("e")
            for c in arg_expr:
                _append_to(e, c)
            acc.append(e)
            result.append(acc)
            i = consumed + 1
            continue
        elif tok in ("sin", "cos", "tan", "log", "ln", "exp",
                     "lim", "liminf", "limsup", "max", "min", "sup", "inf", "det",
                     "dim", "gcd", "ker", "Pr", "argmax", "argmin", "arg",
                     "arcsin", "arccos", "arctan", "sinh", "cosh", "tanh", "coth",
                     "csc", "sec", "cot", "deg", "lg", "hom",
                     "sgn", "diag", "tr", "trace", "rank", "span", "erf", "erfc", "adj",
                     "injlim", "projlim", "varinjlim", "varprojlim", "varliminf", "varlimsup") and is_cmd:
            result.append(_omml_run(tok, normal=True))
            i += 1
            continue
        elif tok in ("big", "Big", "bigg", "Bigg") and is_cmd:
            # \big( \big[ \Big( etc — size hint, OMML auto-stretches delimiters; skip
            i += 1
            continue
        elif tok == "prod" and is_cmd:
            nary = _omml_elem("nary")
            nary_pr = _omml_elem("naryPr")
            chr_el = _omml_elem("chr")
            chr_el.set(qn("m:val"), "∏")
            nary_pr.append(chr_el)
            nary.append(nary_pr)
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "_":
                sub_expr, sub_consumed = _parse_arg(tokens, i + 2)
                sub_el = _omml_elem("sub")
                for c in sub_expr:
                    _append_to(sub_el, c)
                nary.append(sub_el)
                i = sub_consumed
            if i + 1 < len(tokens) and _tok(tokens, i + 1) == "^":
                sup_expr, sup_consumed = _parse_arg(tokens, i + 2)
                sup_el = _omml_elem("sup")
                for c in sup_expr:
                    _append_to(sup_el, c)
                nary.append(sup_el)
                i = max(i, sup_consumed)
            result.append(nary)
            i += 1
            continue
        elif tok in ("{", "}"):
            # Literal braces from \{ \} — render as stretchy delimiters
            d = _omml_elem("d")
            d_pr = _omml_elem("dPr")
            beg = _omml_elem("begChr")
            beg.set(qn("m:val"), "{")
            d_pr.append(beg)
            end = _omml_elem("endChr")
            end.set(qn("m:val"), "}")
            d_pr.append(end)
            d.append(d_pr)
            e = _omml_elem("e")
            inner_tokens = []
            depth = 1
            j = i + 1
            while j < len(tokens) and depth > 0:
                if _tok(tokens, j) == "{":
                    depth += 1
                elif _tok(tokens, j) == "}":
                    depth -= 1
                    if depth == 0:
                        break
                inner_tokens.append(tokens[j])
                j += 1
            inner = _parse_expr(inner_tokens)
            if inner:
                for c in inner:
                    _append_to(e, c)
            d.append(e)
            result.append(d)
            i = j + 1
            continue
        else:
            # Plain text or command symbol lookup
            text = _DOC_SYMBOLS.get(tok, tok) if is_cmd else tok
            # Structural punctuation/digits/delimiters → upright (not italic)
            _UPRIGHT_CHARS = set("()[]{}.,;:=0123456789&")
            upright = all(c in _UPRIGHT_CHARS for c in tok)
            result.append(_omml_run(text, normal=upright))
            i += 1
            continue
        i += 1
    return result


def _tok(tokens, idx):
    """Get the raw token string at index idx (strips is_command flag)."""
    return tokens[idx][0] if idx < len(tokens) else ""


def _parse_arg(tokens: list, start: int) -> tuple[list, int]:
    """Parse a single argument: either a single token or { ... } group."""
    if start >= len(tokens):
        return ([], start)
    if _tok(tokens, start) == "{":
        depth = 1
        inner = []
        i = start + 1
        while i < len(tokens) and depth > 0:
            t = _tok(tokens, i)
            if t == "{":
                depth += 1
            elif t == "}":
                depth -= 1
                if depth == 0:
                    break
            inner.append(tokens[i])
            i += 1
        return (_parse_expr(inner), i)
    else:
        text, is_cmd = tokens[start]
        mapped = _DOC_SYMBOLS.get(text, text) if is_cmd else text
        return ([_omml_run(mapped)], start)


def _tokenize(latex: str, pos: int) -> tuple[list, int]:
    """Tokenize a LaTeX expression into (token, is_command) tuples.

    is_command=True means the token came from a \\command sequence and should
    be eligible for structural handling (\\frac, \\sum, \\int, \\mathbf, etc.).
    Plain text tokens (is_command=False) bypass command handlers to prevent
    conflicts like `int` in `^{int}` being misinterpreted as the \\int command.
    """
    _SPECIAL = "\\^_{} \t\n\r,;()[]|&"
    tokens = []
    i = pos
    while i < len(latex):
        ch = latex[i]
        if ch == "\\":
            j = i + 1
            # Handle \| (norm/double-bar)
            if j < len(latex) and latex[j] == "|":
                tokens.append(("Vert", True))
                i = j + 1
                continue
            # Handle \\ (matrix row separator)
            if j < len(latex) and latex[j] == "\\":
                tokens.append(("\\", True))
                i = j + 1
                continue
            while j < len(latex) and (latex[j].isalpha() or latex[j] == "*"):
                j += 1
            cmd = latex[i + 1:j]
            if cmd:
                tokens.append((cmd, True))  # Mark as command
            i = j
        elif ch in _SPECIAL:
            if not ch.isspace():
                tokens.append((ch, False))  # Mark as literal punctuation
            i += 1
        else:
            j = i
            while j < len(latex) and latex[j] not in _SPECIAL:
                j += 1
            tokens.append((latex[i:j], False))  # Plain text
            i = j
    return (tokens, i)
    return (tokens, i)


def _split_matrix_rows(tokens: list) -> list:
    """Split matrix tokens into rows (by \\) and columns (by &)."""
    rows = []
    current_row = []
    current_cell = []
    i = 0
    while i < len(tokens):
        t, is_cmd = tokens[i]
        if is_cmd and t == "\\":
            # Row separator
            current_row.append(current_cell)
            current_cell = []
            rows.append(current_row)
            current_row = []
            i += 1
        elif t == "&" and not is_cmd:
            # Column separator
            current_row.append(current_cell)
            current_cell = []
            i += 1
        else:
            current_cell.append(tokens[i])
            i += 1
    # Don't forget the last cell and row
    if current_cell or current_row:
        current_row.append(current_cell)
    if current_row:
        rows.append(current_row)
    return rows


def _delim_char(ch: str) -> str:
    """Map simple delimiter chars to OMML delimiter character values."""
    mapping = {"(": "(", ")": ")", "[": "[", "]": "]",
               "{": "{", "}": "}", "|": "|", ".": ""}
    return mapping.get(ch, ch)


def _append_to(parent: OxmlElement, child):
    """Append a child, wrapping plain OxmlElements or unwrapping lists."""
    if isinstance(child, list):
        for c in child:
            parent.append(c)
    else:
        parent.append(child)


def _insert_omml_paragraph(paragraph, omath: OxmlElement):
    """Insert an OMML equation into a paragraph as inline math."""
    run = OxmlElement("w:r")
    run.append(omath)
    paragraph._element.append(run)


def _insert_display_omml(doc, omath: OxmlElement):
    """Insert a display-style OMML equation as a centered paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = OxmlElement("m:oMathPara")
    omath_para.append(omath)
    run = OxmlElement("w:r")
    run.append(omath_para)
    para._element.append(run)
    return para


# ── Style setup ────────────────────────────────────────────────────

def _setup_styles(doc: Document):
    """Configure document styles for CJK patent disclosure."""
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(12)  # 小四
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 2.0  # 2× for patent readability

    # Heading styles
    for level, (size, font_heading) in enumerate([
        (SIZE_H1, True), (SIZE_H2, True), (SIZE_H3, True)
    ], start=1):
        hs = doc.styles[f"Heading {level}"]
        hs.font.size = Pt(size / 2)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


# ── Markdown parsing ───────────────────────────────────────────────

def parse_markdown_sections(content: str) -> list[dict]:
    """Parse markdown content into structured sections."""
    sections = []
    current_section = None

    for line in content.split("\n"):
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        img_match = re.match(r"!\[(.*?)\]\((.*?)\)", line.lstrip())

        if heading_match:
            if current_section:
                sections.append(current_section)
            current_section = {
                "level": len(heading_match.group(1)),
                "title": heading_match.group(2).strip(),
                "lines": [],
                "images": [],
            }
        elif img_match and current_section:
            current_section["images"].append(
                {"alt": img_match.group(1), "path": img_match.group(2)}
            )
        elif current_section is not None:
            current_section["lines"].append(line)

    if current_section:
        sections.append(current_section)
    return sections


def _parse_inline_formatting(paragraph, text: str, font_heading: bool = False):
    """Parse inline Markdown + LaTeX and add formatted runs.

    Uses sequential scanning (not re.split) to avoid duplicate capture-group
    output that causes formula text to appear twice.
    """
    # Combined pattern: **bold** | *italic* | $$display$$ | $inline$
    pattern = re.compile(
        r"\*\*(.+?)\*\*"          # **bold**  (group 1)
        r"|\*(.+?)\*"             # *italic*  (group 2)
        r"|\$\$([^$]+?)\$\$"      # $$display$$ (group 3)
        r"|\$([^$]+?)\$"          # $inline$   (group 4)
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        plain = text[last_end:m.start()]
        if plain:
            paragraph.add_run(plain)

        if m.group(1):  # **bold**
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif m.group(2):  # *italic*
            run = paragraph.add_run(m.group(2))
            run.italic = True
        elif m.group(3) or m.group(4):  # $$display$$ or $inline$
            latex = (m.group(3) or m.group(4)).strip()
            omath = _latex_to_omml(latex)
            if omath is not None:
                _insert_omml_paragraph(paragraph, omath)
            else:
                run = paragraph.add_run(f" {latex} ")
                run.font.size = Pt(10)
                run.font.name = "Consolas"
                run.font.color.rgb = RGBColor(80, 80, 80)

        last_end = m.end()

    # Remaining plain text
    plain = text[last_end:]
    if plain:
        paragraph.add_run(plain)

    # Apply CJK fonts to all text runs (skip OMML elements)
    for run in paragraph.runs:
        if run._element.find(qn('m:oMath')) is None:
            _set_cjk_font(run, font_heading=font_heading)


def _add_paragraph(doc, text: str, font_heading: bool = False):
    """Add a body paragraph with full formatting."""
    # Strip markdown list markers (*, -, +) from the start of lines
    cleaned = re.sub(r"^[\*\-+]\s+", "", text, flags=re.MULTILINE)
    para = doc.add_paragraph()
    _set_line_spacing(para)
    _parse_inline_formatting(para, cleaned, font_heading=font_heading)
    return para


def _is_table_row(line: str) -> bool:
    """Check if a line is a markdown table row."""
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def _is_table_separator(line: str) -> bool:
    """Check if a line is a markdown table separator (e.g. |---|:--:|---|)."""
    s = line.strip()
    if not (s.startswith("|") and s.endswith("|")):
        return False
    inner = s[1:-1]
    return all(c in "|-: " for c in inner)


def _parse_table_row(line: str) -> list:
    """Parse a markdown table row into cell contents."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc, rows: list, headers: list = None):
    """Add a formatted table with inline math and CJK styling.

    Design inspired by minimax-docx typography guide:
    - Header: bold white text on blue background
    - Body: alternating row shading for readability
    - Cell text processed through inline formatting (math → OMML, bold, etc.)
    """
    if headers:
        all_rows = [headers] + rows
    else:
        all_rows = rows
    ncols = max(len(r) for r in all_rows)
    nrows = len(all_rows)

    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    # Set preferred width to 100%
    tbl_pr = table._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), '5000')
    tbl_w.set(qn('w:type'), 'pct')
    tbl_pr.append(tbl_w)

    for r_idx, row_data in enumerate(all_rows):
        is_header = headers is not None and r_idx == 0
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= ncols:
                break
            cell = table.cell(r_idx, c_idx)
            # Clear default paragraph and add formatted content
            para = cell.paragraphs[0]
            para.clear()
            _set_line_spacing(para, 276)  # 1.15× for table cells

            # Process inline math ($...$), bold, italic in cell text
            _parse_inline_formatting(para, cell_text, font_heading=is_header)

            # Style all runs in the cell
            for run in para.runs:
                run.font.size = Pt(9)
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

            # Cell shading
            shading = None
            if is_header:
                # Blue header background
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="2563EB" w:val="clear"/>'
                )
            elif r_idx % 2 == 0:
                # Alternating light gray for body rows
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:val="clear"/>'
                )
            if shading is not None:
                cell._element.get_or_add_tcPr().append(shading)

            # Center-align header cells, left-align body cells
            if is_header:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing after table


def _add_heading_section(doc, title: str, level: int):
    """Add a heading with CJK font."""
    heading = doc.add_heading(title, level=min(level, 3))
    for run in heading.runs:
        _set_cjk_font(run, font_heading=True)
    return heading


# ── Main conversion ────────────────────────────────────────────────

def markdown_to_docx(md_path: Path, output_path: Path) -> None:
    """Convert a Markdown disclosure file to a .docx document."""
    global _DOC_SYMBOLS, _UNKNOWN_COMMANDS
    content = md_path.read_text(encoding="utf-8")

    # ── Step 0: Build per-document symbol map ──
    _DOC_SYMBOLS, _UNKNOWN_COMMANDS = _build_symbol_map(content)
    if _UNKNOWN_COMMANDS:
        print(f"Unknown LaTeX commands ({len(_UNKNOWN_COMMANDS)}): {', '.join(sorted(_UNKNOWN_COMMANDS))}")
        # Save unknown commands to a reference file
        report_path = output_path.parent / f"{md_path.stem}_unknown_symbols.json"
        report_path.write_text(
            json.dumps(sorted(_UNKNOWN_COMMANDS), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"  → Saved to: {report_path}")
        print(f"  → These will render as literal text in the DOCX. Map them to Unicode in the symbol library to fix.")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_TOP)
        section.bottom_margin = Cm(MARGIN_BOTTOM)
        section.left_margin = Cm(MARGIN_LEFT)
        section.right_margin = Cm(MARGIN_RIGHT)

    _setup_styles(doc)

    # Document title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("专利交底书")
    title_run.bold = True
    title_run.font.size = Pt(SIZE_TITLE / 2)
    _set_cjk_font(title_run, font_heading=True)

    sections = parse_markdown_sections(content)

    for sec in sections:
        # Skip the markdown H1 title (it's usually the doc title)
        if sec["level"] == 1 and sec["title"].startswith("#"):
            continue

        _add_heading_section(doc, sec["title"], sec["level"])

        # Body text — process tables and paragraphs in one pass
        sec_lines = sec["lines"]
        if sec_lines:
            # Build groups: table blocks and paragraph blocks
            i = 0
            while i < len(sec_lines):
                line = sec_lines[i]
                # Check for table: header row followed by separator row
                if (i + 1 < len(sec_lines) and _is_table_row(line)
                        and _is_table_separator(sec_lines[i + 1])):
                    headers = _parse_table_row(line)
                    data_rows = []
                    j = i + 2
                    while j < len(sec_lines) and _is_table_row(sec_lines[j]):
                        data_rows.append(_parse_table_row(sec_lines[j]))
                        j += 1
                    _add_table(doc, data_rows, headers)
                    i = j
                    continue
                # Regular paragraph: collect until blank line or next table
                para_lines = []
                while i < len(sec_lines) and sec_lines[i].strip() and not (
                    _is_table_row(sec_lines[i])
                    and i + 1 < len(sec_lines)
                    and _is_table_separator(sec_lines[i + 1])
                ):
                    para_lines.append(sec_lines[i])
                    i += 1
                if para_lines:
                    text = "\n".join(para_lines).strip()
                    if text:
                        _add_paragraph(doc, text)
                # Skip blank lines
                while i < len(sec_lines) and not sec_lines[i].strip():
                    i += 1

        # Images — resolve relative paths against the markdown file's directory
        md_dir = md_path.parent
        for img in sec["images"]:
            img_path = Path(img["path"])
            if not img_path.is_absolute():
                img_path = md_dir / img_path
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(img["alt"])
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(100, 100, 100)
                _set_cjk_font(cap_run, font_heading=False)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Convert patent disclosure Markdown to Word (.docx)"
    )
    parser.add_argument("input", help="Path to input Markdown file")
    parser.add_argument("-o", "--output", default="disclosure.docx",
                        help="Output .docx path")
    args = parser.parse_args()

    markdown_to_docx(Path(args.input), Path(args.output))


if __name__ == "__main__":
    main()
